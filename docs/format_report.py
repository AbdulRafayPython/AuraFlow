"""
AuraFlow FYP Report – Professional Formatter
=============================================
1. Removes the old (first) Chapter 4 (body indices 169-253).
2. Moves Chapter 6 AFTER the new Chapter 4.
3. Applies a consistent professional theme throughout.
"""

import copy, re
from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT

# ── Colour palette ──────────────────────────────────────────────────────
NAVY      = RGBColor(0x1F, 0x3A, 0x6E)   # headings
MID_BLUE  = RGBColor(0x2B, 0x5C, 0x99)   # sub-headings / table headers
LIGHT_BG  = RGBColor(0xEB, 0xF2, 0xF7)   # table alt rows
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_CODE = RGBColor(0xF4, 0xF4, 0xF4)
BORDER_C  = RGBColor(0x1F, 0x3A, 0x6E)

# ── Helpers ──────────────────────────────────────────────────────────────

def hex_color(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(rgb))
    old = tcPr.find(qn('w:shd'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """kwargs: top/bottom/left/right with (sz, color) tuples."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, (sz, color) in kwargs.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(tcBorders)


def set_table_border(table):
    """Outer and inner border."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    border_color = hex_color(BORDER_C)
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), border_color)
        tblBorders.append(el)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)


def set_table_width(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    old = tblPr.find(qn('w:tblW'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblW)


def para_xml_text(p_elem) -> str:
    return ''.join(t.text or '' for t in p_elem.iter(qn('w:t'))).strip()


def clean_heading_text(text: str) -> str:
    """Remove trailing/leading dashes and fix numbered section spacing."""
    text = re.sub(r'\s*[\u2013\u2014-]+\s*', ' ', text)   # em/en/hyphen dash
    # Only remove colon-space in numbered section headings (e.g. "6.1 : Title")
    text = re.sub(r'^(\d+\.\d+(?:\.\d+)?)\s*:\s+', r'\1 ', text)
    # Remove trailing colon from chapter titles ("Chapter 6:" → "Chapter 6")
    text = re.sub(r'^(Chapter\s+\d+):\s*$', r'\1', text)
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()


# ── Style helpers ────────────────────────────────────────────────────────

def _clear_para_direct_formatting(para):
    """Remove per-run fonts/colors so the paragraph style takes precedence."""
    for run in para.runs:
        run.font.color.rgb = None
        run.font.name = None
        run.font.size = None
        run.bold = None
        run.italic = None


def apply_normal_style(para, doc):
    para.style = doc.styles['Normal']
    fmt = para.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.15
    fmt.first_line_indent = None
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
        run.bold = None
        run.italic = None


def apply_chapter_style(para, doc):
    """Big chapter title: CHAPTER X, centered, Navy, 20pt Bold."""
    para.style = doc.styles['Heading 1']
    fmt = para.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt.space_before = Pt(18)
    fmt.space_after  = Pt(6)
    fmt.page_break_before = True
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(20)
        run.font.color.rgb = NAVY
        run.bold = True
        run.italic = False
    # If no runs, set via XML
    if not para.runs:
        r = para.add_run(para.text)
        r.font.name = 'Calibri'
        r.font.size = Pt(20)
        r.font.color.rgb = NAVY
        r.bold = True


def apply_h2_style(para, doc):
    """Section heading: 14pt Bold Navy."""
    para.style = doc.styles['Heading 2']
    fmt = para.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt.space_before = Pt(14)
    fmt.space_after  = Pt(4)
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.color.rgb = NAVY
        run.bold = True
        run.italic = False


def apply_h3_style(para, doc):
    """Sub-section heading: 12pt Bold Mid-Blue."""
    para.style = doc.styles['Heading 3']
    fmt = para.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt.space_before = Pt(10)
    fmt.space_after  = Pt(3)
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = MID_BLUE
        run.bold = True
        run.italic = False


def apply_code_style(para, doc):
    """Monospaced code block."""
    para.style = doc.styles['Normal']
    fmt = para.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(0)
    fmt.left_indent  = Cm(1)
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        run.bold = False


def apply_caption_style(para, doc):
    """Figure/table caption: italic, centered, 10pt."""
    para.style = doc.styles['Normal']
    fmt = para.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt.space_before = Pt(2)
    fmt.space_after  = Pt(8)
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        run.italic = True
        run.bold = False


# ── Table formatter ───────────────────────────────────────────────────────

def format_table(table):
    set_table_width(table)
    set_table_border(table)
    border_hex = hex_color(BORDER_C)

    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)
        is_alt    = (row_idx % 2 == 0 and not is_header)

        for cell in row.cells:
            # Background
            if is_header:
                set_cell_bg(cell, NAVY)
            elif is_alt:
                set_cell_bg(cell, LIGHT_BG)
            else:
                set_cell_bg(cell, WHITE)

            # Cell borders
            border_args = {side: ('4', border_hex)
                           for side in ('top','bottom','left','right')}
            set_cell_border(cell, **border_args)

            # Cell vertical alignment
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            old_va = tcPr.find(qn('w:vAlign'))
            if old_va is not None:
                tcPr.remove(old_va)
            tcPr.append(vAlign)

            # Text formatting inside cell
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after  = Pt(3)
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    if is_header:
                        run.font.color.rgb = WHITE
                        run.bold = True
                    else:
                        run.font.color.rgb = BLACK
                        run.bold = False


# ── Section classifier ────────────────────────────────────────────────────

# Known subsection labels (H3)
H3_LABELS = {
    'frontend implementation', 'backend implementation', 'ai agents implementation',
    'ai agent integration', 'backend architecture', 'frontend technologies',
    'backend technologies', 'ai agents and libraries', 'databases and storage',
    'communication and media tools', 'development tools',
    'voice communication', 'code snippet',
    'unit testing', 'integration testing', 'system testing',
    'user acceptance testing (uat)',
    'ai agent benchmark', 'system performance evaluation',
    'use case diagram', 'system architecture diagram',
    'entity relationship (er) diagram', 'activity diagram',
    'sequence diagram', 'class diagram',
    # Chapter 1 sub-labels
    'core goals:', 'intelligent features:',
    # Requirement analysis sub-sections (numbered)
    '1. user account management', '2. community and channel management',
    '3. real time messaging', '4. ai powered agents',
    '5. voice communication', '6. admin dashboard',
}

# Front matter headings
FRONT_MATTER = {'abstract', 'acknowledgement', 'acknowledgements', 'dedication'}

# Patterns
# A chapter heading starts with "Chapter N", is short, and is NOT a sentence
RE_CHAPTER   = re.compile(r'^Chapter\s+\d+', re.IGNORECASE)
# Body text that mentions chapters (ToC descriptions)
RE_CHAPTER_BODY = re.compile(
    r'\b(provides|presents|covers|showcases|wraps|demonstrates|explains|outlines)\b',
    re.IGNORECASE
)
RE_SECTION   = re.compile(r'^\d+\.\d+(\s+|$)')
RE_SUBSECT   = re.compile(r'^\d+\.\d+\.\d+(\s+|$)')
RE_CODE_LINE = re.compile(
    r'^(#\s|@|def |from |import |class |    |sio\.|result\s|REDIS|sio =|summarizer|'
    r"'channel_id|moderate_message_task|@celery_app|check_user|'content'|}, room=)"
)
RE_CAPTION   = re.compile(r'^(figure|table)\s+\d+', re.IGNORECASE)
RE_TABLE_REF = re.compile(r'^table\s+\d+[\.:]\s*', re.IGNORECASE)
RE_NUMBLIST  = re.compile(r'^\d+\.\s+\w')


def classify(text: str):
    t = text.strip()
    tl = t.lower()
    if not t:
        return 'empty'
    if tl in FRONT_MATTER:
        return 'front_matter'
    if RE_CHAPTER.match(t) and len(t) <= 70 and not RE_CHAPTER_BODY.search(t):
        return 'chapter'
    if t in ('Implementation & Results', 'System Design & Methodology',
             'Conclusion & Future Work'):
        return 'chapter_sub'
    if RE_SUBSECT.match(t):
        return 'h3'
    if RE_SECTION.match(t):
        return 'h2'
    if tl in H3_LABELS:
        return 'h3'
    if RE_TABLE_REF.match(t):
        return 'caption'
    if RE_CAPTION.match(t):
        return 'caption'
    if RE_CODE_LINE.match(t):
        return 'code'
    # Multi-line code detection: contains only code-like tokens
    if '\n' not in t and t.startswith(("'", '"', '}', '{', '(', ')')):
        return 'code'
    return 'body'


# ── Page setup ────────────────────────────────────────────────────────────

def setup_page(doc):
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)


# ── Define/update built-in styles ─────────────────────────────────────────

def define_styles(doc):
    styles = doc.styles

    def _set(style_name, font_name, font_size, bold, color, align,
             space_before, space_after, line_spacing=1.15):
        try:
            s = styles[style_name]
        except KeyError:
            return
        s.font.name = font_name
        s.font.size = Pt(font_size)
        s.font.bold = bold
        s.font.color.rgb = color
        pf = s.paragraph_format
        pf.alignment = align
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing

    _set('Normal',    'Calibri', 11, False, BLACK, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 6)
    _set('Heading 1', 'Calibri', 20, True,  NAVY,  WD_ALIGN_PARAGRAPH.CENTER,  18, 6)
    _set('Heading 2', 'Calibri', 14, True,  NAVY,  WD_ALIGN_PARAGRAPH.LEFT,    14, 4)
    _set('Heading 3', 'Calibri', 12, True,  MID_BLUE, WD_ALIGN_PARAGRAPH.LEFT, 10, 3)


# ── Main ─────────────────────────────────────────────────────────────────

def main():
    src_path = r'z:\FYP\AuraFlow\docs\FYP-1_FINAL_REPORT_Chapter4_Updated.docx'
    out_path = r'z:\FYP\AuraFlow\docs\FYP-1_FINAL_REPORT_FORMATTED.docx'

    doc = Document(src_path)
    body = doc.element.body
    all_children = list(body)

    # ── STEP 1: Identify ranges ──────────────────────────────────────────
    # OLD Chapter 4: body indices 169–253  (del)
    # Chapter 6:     body indices 254–272  (keep, reorder after new ch4)
    # NEW Chapter 4: body indices 273–316  (keep)
    # sectPr:        body index  317       (keep last)

    # We need to determine boundaries dynamically
    old_ch4_start = old_ch4_end = None
    ch6_start = ch6_end = None
    new_ch4_start = new_ch4_end = None

    for i, child in enumerate(all_children):
        tag = child.tag.split('}')[-1]
        if tag != 'p':
            continue
        text = para_xml_text(child)

        # First "Chapter 4" that appears WITHOUT ":" (old one)
        if old_ch4_start is None and re.match(r'^Chapter 4$', text):
            old_ch4_start = i

        # Chapter 6 standalone heading (not a sentence referencing ch6)
        if ch6_start is None and re.match(r'^Chapter 6[:\s]*$', text):
            ch6_start = i

        # New Chapter 4 (has colon: "Chapter 4: Implementation")
        if new_ch4_start is None and re.match(r'^Chapter 4:', text):
            new_ch4_start = i

    # old_ch4_end = just before ch6_start
    old_ch4_end = ch6_start - 1

    # ch6_end = just before new_ch4_start (skip blanks)
    ch6_end = new_ch4_start - 1

    # new_ch4_end = last non-sectPr element
    new_ch4_end = len(all_children) - 1
    while all_children[new_ch4_end].tag.split('}')[-1] == 'sectPr':
        new_ch4_end -= 1

    print(f"Old Ch4: [{old_ch4_start}–{old_ch4_end}]")
    print(f"Ch6:     [{ch6_start}–{ch6_end}]")
    print(f"New Ch4: [{new_ch4_start}–{new_ch4_end}]")

    # ── STEP 2: Build new body order ─────────────────────────────────────
    # Segments:
    # A: front matter + ch1 + ch2 + ch3 → indices 0 .. old_ch4_start-1
    # B: new Ch4                          → new_ch4_start .. new_ch4_end
    # C: Ch6                              → ch6_start .. ch6_end
    # D: sectPr

    seg_A = all_children[:old_ch4_start]
    seg_B = all_children[new_ch4_start:new_ch4_end + 1]
    seg_C = all_children[ch6_start:ch6_end + 1]
    sect  = [c for c in all_children if c.tag.split('}')[-1] == 'sectPr']

    # Rebuild body XML
    # Remove all children from body
    for child in list(body):
        body.remove(child)

    # Add back in order: A, blank, B, blank, C, sectPr
    def add_page_break_para(parent):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pb = OxmlElement('w:pageBreak')
        r = OxmlElement('w:r')
        r.append(pb)
        p.append(pPr)
        p.append(r)
        parent.append(p)

    for el in seg_A:
        body.append(el)

    for el in seg_B:
        body.append(el)

    for el in seg_C:
        body.append(el)

    for el in sect:
        body.append(el)

    print("Body rebuilt successfully.")

    # ── STEP 3: Page setup ───────────────────────────────────────────────
    setup_page(doc)

    # ── STEP 4: Define base styles ───────────────────────────────────────
    define_styles(doc)

    # ── STEP 5: Format every paragraph ───────────────────────────────────
    prev_class = None
    for para in doc.paragraphs:
        raw = para.text.strip()
        if not raw:
            # Remove excessive blank paragraphs (keep max 1)
            continue

        # Clean text in runs (remove lone dashes from headings)
        kind = classify(raw)

        if kind == 'front_matter':
            para.style = doc.styles['Heading 1']
            fmt = para.paragraph_format
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_before = Pt(24)
            fmt.space_after  = Pt(12)
            fmt.page_break_before = False
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(16)
                run.font.color.rgb = NAVY
                run.bold = True
                run.italic = False

        elif kind == 'chapter':
            # Clean text
            clean = clean_heading_text(raw)
            # Rebuild runs
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = clean
            apply_chapter_style(para, doc)

        elif kind == 'chapter_sub':
            # Subtitle line below chapter number (same heading block)
            clean = clean_heading_text(raw)
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = clean
            para.style = doc.styles['Heading 1']
            para.paragraph_format.page_break_before = False
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(10)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(16)
                run.font.color.rgb = NAVY
                run.bold = True

        elif kind == 'h2':
            clean = clean_heading_text(raw)
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = clean
            apply_h2_style(para, doc)

        elif kind == 'h3':
            clean = clean_heading_text(raw)
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = clean
            apply_h3_style(para, doc)

        elif kind == 'caption':
            apply_caption_style(para, doc)

        elif kind == 'code':
            apply_code_style(para, doc)

        else:  # body
            apply_normal_style(para, doc)

        prev_class = kind

    # ── STEP 6: Format all tables ─────────────────────────────────────────
    for table in doc.tables:
        format_table(table)

    # ── STEP 7: Fix Chapter/Section title text cleanup ────────────────────
    # Extra pass: fix chapter heading texts that span multiple paragraphs
    # e.g., "Chapter 1 – Introduction" stored as one paragraph
    for para in doc.paragraphs:
        raw = para.text.strip()
        if not raw:
            continue
        if para.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
            clean = clean_heading_text(raw)
            if clean != raw:
                # Rewrite text preserving run structure
                all_text = ''.join(r.text for r in para.runs)
                if all_text.strip():
                    clean_all = clean_heading_text(all_text)
                    # Clear all runs and set on first run
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = clean_all

    doc.save(out_path)
    print(f"\nFormatted document saved to:\n  {out_path}")


if __name__ == '__main__':
    main()
